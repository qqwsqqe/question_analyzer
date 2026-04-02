"""Word文档提取器"""

import logging
from pathlib import Path
from docx import Document

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class WordExtractor:
    """Word文档文本提取器"""

    def extract(self, file_path: str) -> str:
        """
        提取Word文档文本内容

        Args:
            file_path: Word文件路径(.docx)

        Returns:
            提取的文本内容
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        if file_path.suffix.lower() not in ['.docx']:
            raise ValueError(f"不支持的文件格式: {file_path.suffix}，仅支持.docx")

        text_parts = []

        try:
            doc = Document(file_path)

            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # 提取表格中的文字
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_texts.append(cell.text.strip())
                    if row_texts:
                        text_parts.append(' | '.join(row_texts))

            logger.info(f"成功提取Word文档: {file_path.name}")

        except Exception as e:
            logger.error(f"Word提取失败: {e}")
            raise

        return '\n\n'.join(text_parts)

    def extract_with_structure(self, file_path: str) -> dict:
        """
        提取Word文档并保留结构信息

        Args:
            file_path: Word文件路径

        Returns:
            包含段落和表格的字典
        """
        file_path = Path(file_path)
        doc = Document(file_path)

        result = {
            'paragraphs': [],
            'tables': []
        }

        # 提取带样式的段落
        for para in doc.paragraphs:
            if para.text.strip():
                result['paragraphs'].append({
                    'text': para.text,
                    'style': para.style.name if para.style else 'Normal',
                    'is_bold': any(run.bold for run in para.runs),
                    'is_heading': para.style.name.startswith('Heading') if para.style else False
                })

        # 提取表格
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            result['tables'].append(table_data)

        return result
